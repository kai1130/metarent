from docx import Document
import docx2txt
import openai
import keys
import dateutil.parser as dparser
import re
import datetime

openai.api_key = keys.openai_api_key

def extract_text(docpath):
    doc = docx2txt.process(docpath)
    doc = re.sub(r'\n{2,2}', r"\n", doc)
    doc = re.sub(r'\n{3,}', r"\n\n", doc)
    return doc

def generate_response(task, doc):
    response = openai.Completion.create(
        model="text-davinci-002",
        prompt=f"{task}\n\n{doc}",
        temperature=0,
        max_tokens=500,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0
    )
    return response

def clean_response(response):
    
    def dictify_response(response):
        res = response['choices'][0]['text']
        res = re.sub(r'\n', r'', res)
        res = eval(res)
        return res

    def parse_str_to_num(response_dict):

        currency_fields = ['Base Fee','Fuel Fee','Security Deposit','Rental Term']
        for field in currency_fields:
            response_dict[field] = int(''.join(filter(lambda i: i.isdigit(), response_dict[field])))
        return response_dict

    def parse_datetime(response):

        date_fields = ['Start Date']
        for field in date_fields:
            response[field] = dparser.parse(response[field],fuzzy=True).date()
            
        start_time = datetime.datetime.combine(response['Start Date'], datetime.datetime.min.time())
        response['Start Timestamp'] = int(datetime.datetime.timestamp(start_time)*1000)
        
        return response

    def parse_nums(response):

        num_fields = ['Base Fee', 'Fuel Fee', 'Security Deposit','Rental Term']
        for field in num_fields:
            response[field] = int(response[field])
        return response

    def clean_strs(response):

        for k, v in response.items():
            if isinstance(v, str):
                response[k] = v.strip()
        return response

    response = dictify_response(response)
    response = parse_str_to_num(response)
    response = parse_datetime(response)
    response = parse_nums(response)
    response = clean_strs(response)

    return response

def validate_response(response):
    
    validation_list = []
    
    def validate_email(email, party):
        pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        if re.match(pattern, email):
            validation_list.append(True)
            print(f"    Valid - {party} Email")
        else:
            validation_list.append(False)
            print(f"    Invalid - {party} Email")
        
    def validate_name(name, party):
        if name.replace(" ", "").isalpha() and (bool(re.search('\s', name))):
            validation_list.append(True)
            print(f"    Valid - {party} Name")
        else:
            validation_list.append(False)
            print(f"    Invalid - {party} Name")
    
    def validate_party(party, response):
        print(f"Validating {party}")
        
        validate_email(response[f'{party} Email'], party)
        validate_name(response[f'{party} Name'], party)
        
    def validate_car(response):
        print(f"Validating Car")
        
        for field in ['Car Make','Car Model','Car VIN','Car Year','Car Color']:
            if response[field]:
                validation_list.append(True)
                print(f"    Valid - {field}")
            else:
                validation_list.append(False)
                print(f"    Invalid - {field}")
    
    def validate_dates(response):
        print(f"Validating Dates")
        
        if (type(response["Start Date"]) is datetime.date):
            if (response["Start Date"] >= datetime.datetime.now().date()):
                validation_list.append(True)
                print(f"    Valid - Start Date")
                print(f"    Valid - Start Timestamp")
            else:
                validation_list.append(False)
                print(f"    Invalid - Start Date")
                print(f"    Invalid - Start Timestamp")
        else:
            validation_list.append(False)
            print(f"    Invalid - Start Date")
            print(f"    Invalid - Start Timestamp")
    
    def validate_fee(fee, field):
        
        if (type(fee) == int) or (type(fee) == float):
            if fee > 0:
                validation_list.append(True)
                print(f"    Valid - {field}")
            else:
                validation_list.append(False)
                print(f"    Invalid - {field}")
        else:
            validation_list.append(False)
            print(f"    Invalid - {field}")
            
     
    def validate_feeStructure(response):
        print("Validating Fee Structure")
        
        [validate_fee(v, k) for k,v in response.items() if k in ['Base Fee','Fuel Fee','Security Deposit']]
                            
    for party in ['Owner','Renter']:
        validate_party(party, response)
    validate_car(response)
    validate_dates(response)
    validate_feeStructure(response)
    
    if all(validation_list):
        print("**Response Validation Token Set: TRUE**")
        response['Validation Token'] = True
    else:
        print("**Response Validation Token Set: FALSE**")
        response['Validation Token'] = False
    
    return response

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)
        
def replace_doc_text(template_file_path, output_file_path, variables):

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)
    print(f"Saved output file to {output_file_path}")